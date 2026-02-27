"""Abstract interface for document loader implementations."""

import datetime
import logging
import mimetypes
import os
import re
import tempfile
import traceback
import uuid
from urllib.parse import urlparse
from xml.etree import ElementTree
from pathlib import Path

import requests
from docx import Document as DocxDocument

# from configs import dify_config
# from core.helper import ssrf_proxy
from core.rag.extractor.extractor_base import BaseExtractor
from core.rag.entities.document import Document
from core.rag.extractor.extractor_utils import get_content_from_url, save_image
from config.config import get_config
from logger import get_logger
# from extensions.ext_database import db
# from extensions.ext_storage import storage
# from models.enums import CreatedByRole
# from models.model import UploadFile

logger = get_logger('rag')


class WordExtractor(BaseExtractor):
    """Load docx files.

    Args:
        file_path: Path to the file to load.
    """

    def __init__(self, file_path: str, file_id: str):
        """Initialize with file path."""
        
        self.file_path_id = os.path.splitext(os.path.basename(file_path))[0]

        self.file_url = get_config().get('dependent_info').get('knowledge_base').get('file_url')
        self.pic_save_path_prefix = get_config().get('dependent_info').get('knowledge_base').get('pic_save_path_prefix')
        self.pic_url_prefix = get_config().get('dependent_info').get('knowledge_base').get('pic_url_prefix')
        if self.file_url.endswith('/'):
            self.file_url = self.file_url[:-1]
        if self.pic_url_prefix.endswith('/'):
            self.pic_url_prefix = self.pic_url_prefix[:-1]
        self.pic_save_path_prefix = Path(self.pic_save_path_prefix) / self.file_path_id
        os.makedirs(self.pic_save_path_prefix, exist_ok=True)

        self.file_path = file_path
        self.file_id = file_id

        # if "~" in self.file_path:
        #     self.file_path = os.path.expanduser(self.file_path)

        # # If the file is a web path, download it to a temporary file, and use that
        # if not os.path.isfile(self.file_path) and self._is_valid_url(self.file_path):
        #     r = requests.get(self.file_path)

        #     if r.status_code != 200:
        #         raise ValueError(f"Check the url of your file; returned status code {r.status_code}")

        #     self.web_path = self.file_path
        #     # TODO: use a better way to handle the file
        #     self.temp_file = tempfile.NamedTemporaryFile()  # noqa SIM115
        #     self.temp_file.write(r.content)
        #     self.file_path = self.temp_file.name
        # elif not os.path.isfile(self.file_path):
        #     raise ValueError(f"File path {self.file_path} is not a valid file or url")

    def __del__(self) -> None:
        if hasattr(self, "temp_file"):
            self.temp_file.close()

    def extract(self) -> list[Document]:
        """Load given path as single page."""
        content = self.parse_docx(self.file_path, "tmp_file")
        return [
            Document(
                page_content=content,
                metadata={"source": self.file_path},
            )
        ]

    @staticmethod
    def _is_valid_url(url: str) -> bool:
        """Check if the url is valid."""
        parsed = urlparse(url)
        return bool(parsed.netloc) and bool(parsed.scheme)

    def _extract_images_from_docx(self, doc, image_folder):
        # os.makedirs(image_folder, exist_ok=True)
        image_count = 0
        image_map = {}

        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1
                if rel.is_external:
                    url = rel.target_ref
                    try:
                        response = get_content_from_url(url)
                    except Exception as e:
                        logger.warning(f'Fail to download "{url}" for "{traceback.format_exc()}"')
                        continue
                    if response.status_code == 200:
                        image_ext = mimetypes.guess_extension(response.headers["Content-Type"])
                        if image_ext is None:
                            continue
                        # file_uuid = str(uuid.uuid4())
                        # file_key = "image_files/" + self.tenant_id + "/" + file_uuid + "." + image_ext
                        # mime_type, _ = mimetypes.guess_type(file_key)

                        image_id = str(uuid.uuid4())
                        file_pure_name = image_id + '.' + image_ext
                        file_key = self.pic_save_path_prefix / file_pure_name
                        save_image(str(file_key), response.content)

                        # storage.save(file_key, response.content)
                    else:
                        continue
                else:
                    image_ext = rel.target_ref.split(".")[-1]
                    if image_ext is None:
                        continue
                    # user uuid as file name
                    # file_uuid = str(uuid.uuid4())
                    # file_key = "image_files/" + self.tenant_id + "/" + file_uuid + "." + image_ext
                    # mime_type, _ = mimetypes.guess_type(file_key)

                    image_id = str(uuid.uuid4())
                    file_pure_name = image_id + '.' + image_ext
                    file_key = self.pic_save_path_prefix / file_pure_name
                    save_image(str(file_key), rel.target_part.blob)

                    # storage.save(file_key, rel.target_part.blob)
                # # save file to db
                # upload_file = UploadFile(
                #     tenant_id=self.tenant_id,
                #     storage_type=dify_config.STORAGE_TYPE,
                #     key=file_key,
                #     name=file_key,
                #     size=0,
                #     extension=str(image_ext),
                #     mime_type=mime_type or "",
                #     created_by=self.user_id,
                #     created_by_role=CreatedByRole.ACCOUNT,
                #     created_at=datetime.datetime.now(datetime.UTC).replace(tzinfo=None),
                #     used=True,
                #     used_by=self.user_id,
                #     used_at=datetime.datetime.now(datetime.UTC).replace(tzinfo=None),
                # )

                # db.session.add(upload_file)
                # db.session.commit()

                

                # image_map[rel.target_part] = f"![image]({dify_config.FILES_URL}/files/{upload_file.id}/file-preview)"
                image_map[rel.target_part] = f"![image]({self.file_url}{self.pic_url_prefix}/{self.file_path_id}/{file_pure_name})"

        return image_map

    def _table_to_markdown(self, table, image_map):
        markdown = []
        # calculate the total number of columns
        total_cols = max(len(row.cells) for row in table.rows)

        header_row = table.rows[0]
        headers = self._parse_row(header_row, image_map, total_cols)
        markdown.append("| " + " | ".join(headers) + " |")
        markdown.append("| " + " | ".join(["---"] * total_cols) + " |")

        for row in table.rows[1:]:
            row_cells = self._parse_row(row, image_map, total_cols)
            markdown.append("| " + " | ".join(row_cells) + " |")
        return "\n".join(markdown)

    def _parse_row(self, row, image_map, total_cols):
        # Initialize a row, all of which are empty by default
        row_cells = [""] * total_cols
        col_index = 0
        for cell in row.cells:
            # make sure the col_index is not out of range
            while col_index < total_cols and row_cells[col_index] != "":
                col_index += 1
            # if col_index is out of range the loop is jumped
            if col_index >= total_cols:
                break
            cell_content = self._parse_cell(cell, image_map).strip()
            cell_colspan = cell.grid_span or 1
            for i in range(cell_colspan):
                if col_index + i < total_cols:
                    row_cells[col_index + i] = cell_content if i == 0 else ""
            col_index += cell_colspan
        return row_cells

    def _parse_cell(self, cell, image_map):
        cell_content = []
        for paragraph in cell.paragraphs:
            parsed_paragraph = self._parse_cell_paragraph(paragraph, image_map)
            if parsed_paragraph:
                cell_content.append(parsed_paragraph)
        unique_content = list(dict.fromkeys(cell_content))
        return " ".join(unique_content)

    def _parse_cell_paragraph(self, paragraph, image_map):
        paragraph_content = []
        for run in paragraph.runs:
            if run.element.xpath(".//a:blip"):
                for blip in run.element.xpath(".//a:blip"):
                    image_id = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    if not image_id:
                        continue
                    image_part = paragraph.part.rels[image_id].target_part

                    if image_part in image_map:
                        image_link = image_map[image_part]
                        paragraph_content.append(image_link)
            else:
                paragraph_content.append(run.text)
        return "".join(paragraph_content).strip()

    def _parse_paragraph(self, paragraph, image_map):
        paragraph_content = []
        for run in paragraph.runs:
            if run.element.xpath(".//a:blip"):
                for blip in run.element.xpath(".//a:blip"):
                    embed_id = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    if embed_id:
                        rel_target = run.part.rels[embed_id].target_ref
                        if rel_target in image_map:
                            paragraph_content.append(image_map[rel_target])
            if run.text.strip():
                paragraph_content.append(run.text.strip())
        return " ".join(paragraph_content) if paragraph_content else ""

    def _extract_hyperlink(self, text: str) -> dict:
        hyperlink_pattern = re.compile(r'HYPERLINK\s+"([^"]+)"')
        l_pattern = re.compile(r'\\l\s+"([^"]+)"')
        t_pattern = re.compile(r'\\t\s+"([^"]+)"')
        o_pattern = re.compile(r'\\o\s+"([^"]+)"')

        hyperlink_match = hyperlink_pattern.findall(text)
        l_match = l_pattern.findall(text)
        t_match = t_pattern.findall(text)
        o_match = o_pattern.findall(text)

        output = {
            'link': hyperlink_match[0] if hyperlink_match else '',
            'l': l_match[0] if l_match else '',
            't': t_match[0] if t_match else '',
            'o': o_match[0] if o_match else '',
        }

        return output

    def parse_docx(self, docx_path, image_folder):
        doc = DocxDocument(docx_path)
        # 创建图片存放的根目录
        # os.makedirs(image_folder, exist_ok=True)

        content = []

        # 从文档中提取图片，涉及文件上传，图片替换为内部超链接
        image_map = self._extract_images_from_docx(doc, image_folder)

        hyperlinks_url = None
        url_pattern = re.compile(r"(http://[^\s+]+//|https://[^\s+]+)\"")
        for para in doc.paragraphs:
            for run in para.runs:
                if run.text and hyperlinks_url:
                    result = f"  [{run.text}]({hyperlinks_url})  "
                    run.text = result
                    hyperlinks_url = None
                if "HYPERLINK" in run.element.xml:
                    try:
                        xml = ElementTree.XML(run.element.xml)
                        x_child = [c for c in xml.iter() if c is not None]
                        for x in x_child:
                            if x_child is None:
                                continue
                            if x.tag.endswith("instrText"):
                                if x.text is None:
                                    continue
                                
                                extract_result = self._extract_hyperlink(x.text)
                                if extract_result.get('link'):
                                    hyperlinks_url = extract_result.get('link')
                                    if extract_result.get('l'):
                                        hyperlinks_url = hyperlinks_url + '#' + extract_result.get('l')

                                # for i in url_pattern.findall(x.text):
                                #     loc_match = re.search(r'\\l\s+"([^"]+)"', x.text)
                                #     location = '#' + loc_match.group(1) if loc_match else ''
                                #     hyperlinks_url = str(i) + location
                    except Exception:
                        logger.exception("Failed to parse HYPERLINK xml")

        def parse_paragraph(paragraph):
            paragraph_content = []
            for run in paragraph.runs:
                if hasattr(run.element, "tag") and isinstance(run.element.tag, str) and run.element.tag.endswith("r"):
                    drawing_elements = run.element.findall(
                        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
                    )
                    for drawing in drawing_elements:
                        blip_elements = drawing.findall(
                            ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
                        )
                        for blip in blip_elements:
                            embed_id = blip.get(
                                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                            )
                            if embed_id:
                                image_part = doc.part.related_parts.get(embed_id)
                                if image_part in image_map:
                                    paragraph_content.append(image_map[image_part])
                if run.text.strip():
                    paragraph_content.append(run.text.strip())
            return "".join(paragraph_content) if paragraph_content else ""

        paragraphs = doc.paragraphs.copy()
        tables = doc.tables.copy()
        for element in doc.element.body:
            if hasattr(element, "tag"):
                if isinstance(element.tag, str) and element.tag.endswith("p"):  # paragraph
                    para = paragraphs.pop(0)
                    parsed_paragraph = parse_paragraph(para)
                    if parsed_paragraph.strip():
                        content.append(parsed_paragraph)
                    else:
                        content.append("\n")
                elif isinstance(element.tag, str) and element.tag.endswith("tbl"):  # table
                    table = tables.pop(0)
                    content.append(self._table_to_markdown(table, image_map))
        return "\n".join(content)
